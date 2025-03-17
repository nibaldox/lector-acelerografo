import numpy as np
import matplotlib.pyplot as plt
from datetime import datetime
import os
from pathlib import Path
from io import BytesIO
import base64

class ReportGenerator:
    def __init__(self):
        """
        Inicializa el generador de reportes
        """
        # Crear directorio para reportes si no existe
        self.report_dir = Path("reportes")
        self.report_dir.mkdir(exist_ok=True)
    
    def generate_report(self, data, analysis_results, output_format="pdf"):
        """
        Genera un reporte automático con los resultados del análisis.
        
        Args:
            data (dict): Datos del registro sísmico
            analysis_results (dict): Resultados del análisis
            output_format (str): Formato de salida ('pdf', 'html', 'docx')
            
        Returns:
            str: Ruta al archivo de reporte generado
        """
        # Nombre del archivo basado en la fecha y hora actual
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"reporte_{data['name'].replace(' ', '_')}_{timestamp}"
        
        if output_format == "pdf":
            return self._generate_pdf_report(data, analysis_results, filename)
        elif output_format == "html":
            return self._generate_html_report(data, analysis_results, filename)
        elif output_format == "docx":
            return self._generate_docx_report(data, analysis_results, filename)
        else:
            raise ValueError(f"Formato de salida '{output_format}' no soportado")
    
    def _generate_pdf_report(self, data, analysis_results, filename):
        """
        Genera un reporte en formato PDF
        """
        try:
            from fpdf import FPDF
        except ImportError:
            raise ImportError("La biblioteca 'fpdf' es necesaria para generar reportes PDF. Instálela con 'pip install fpdf'.")
        
        # Crear PDF
        pdf = FPDF()
        pdf.add_page()
        
        # Configuración de fuentes
        pdf.set_font("Arial", "B", 16)
        
        # Título
        pdf.cell(0, 10, "Reporte de Análisis Sísmico", 0, 1, "C")
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"Registro: {data['name']}", 0, 1, "C")
        pdf.cell(0, 10, f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", 0, 1, "C")
        
        # Información del registro
        pdf.ln(10)
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Información del Registro", 0, 1, "L")
        pdf.set_font("Arial", "", 12)
        
        # Metadatos
        if 'metadata' in data:
            for key, value in data['metadata'].items():
                pdf.cell(0, 8, f"{key}: {value}", 0, 1, "L")
        
        # Parámetros del registro
        pdf.ln(5)
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Parámetros del Registro:", 0, 1, "L")
        pdf.set_font("Arial", "", 12)
        
        # Calcular parámetros
        for component in ['N', 'E', 'Z']:
            if f'{component}_aceleracion' in data:
                acc_data = data[f'{component}_aceleracion']
                pga = np.max(np.abs(acc_data))
                pdf.cell(0, 8, f"PGA Componente {component}: {pga:.4f} g", 0, 1, "L")
        
        # Gráficos
        pdf.ln(10)
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Gráficos de Análisis", 0, 1, "L")
        
        # Función para convertir gráfico de matplotlib a imagen para PDF
        def fig_to_img(fig, dpi=200):
            buf = BytesIO()
            fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
            buf.seek(0)
            return buf
        
        # Gráfico de aceleración
        plt.figure(figsize=(10, 6))
        for component, label in zip(['N', 'E', 'Z'], ['Norte-Sur', 'Este-Oeste', 'Vertical']):
            if f'{component}_aceleracion' in data:
                plt.plot(data['time'], data[f'{component}_aceleracion'], label=f'{label}')
        plt.title('Registro de Aceleración')
        plt.xlabel('Tiempo (s)')
        plt.ylabel('Aceleración (g)')
        plt.grid(True)
        plt.legend()
        plt.tight_layout()
        
        # Guardar gráfico en el PDF
        img_buf = fig_to_img(plt.gcf())
        plt.close()
        
        pdf.image(img_buf, x=10, y=None, w=190)
        
        # Si hay resultados de espectro de respuesta
        if 'response_spectrum' in analysis_results:
            pdf.ln(5)
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, "Espectro de Respuesta:", 0, 1, "L")
            
            # Gráfico de espectro de respuesta
            plt.figure(figsize=(10, 6))
            plt.loglog(analysis_results['response_spectrum']['periods'], 
                      analysis_results['response_spectrum']['Sa'], 
                      label='Pseudo-aceleración')
            plt.title('Espectro de Respuesta')
            plt.xlabel('Período (s)')
            plt.ylabel('Sa (g)')
            plt.grid(True, which="both")
            plt.legend()
            plt.tight_layout()
            
            # Guardar gráfico en el PDF
            img_buf = fig_to_img(plt.gcf())
            plt.close()
            
            pdf.image(img_buf, x=10, y=None, w=190)
        
        # Guardar el PDF
        output_path = self.report_dir / f"{filename}.pdf"
        pdf.output(str(output_path))
        return str(output_path)
    
    def _generate_html_report(self, data, analysis_results, filename):
        """
        Genera un reporte en formato HTML
        """
        try:
            import jinja2
            import webbrowser
        except ImportError:
            raise ImportError("La biblioteca 'jinja2' es necesaria para generar reportes HTML. Instálela con 'pip install jinja2'.")
        
        # Crear plantilla HTML básica
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Reporte de Análisis Sísmico</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1, h2 { color: #333366; }
                .container { max-width: 1000px; margin: 0 auto; }
                .info-section { margin: 20px 0; padding: 15px; background-color: #f5f5f5; border-radius: 5px; }
                .graph-section { margin: 30px 0; }
                table { border-collapse: collapse; width: 100%; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>Reporte de Análisis Sísmico</h1>
                <p><strong>Registro:</strong> {{ data_name }}</p>
                <p><strong>Fecha de generación:</strong> {{ generation_date }}</p>
                
                <div class="info-section">
                    <h2>Información del Registro</h2>
                    {% if metadata %}
                    <table>
                        <tr><th>Parámetro</th><th>Valor</th></tr>
                        {% for key, value in metadata.items() %}
                        <tr><td>{{ key }}</td><td>{{ value }}</td></tr>
                        {% endfor %}
                    </table>
                    {% endif %}
                    
                    <h3>Parámetros del Registro</h3>
                    <table>
                        <tr><th>Componente</th><th>PGA (g)</th></tr>
                        {% for comp, label in [('N', 'Norte-Sur'), ('E', 'Este-Oeste'), ('Z', 'Vertical')] %}
                        {% if comp + '_aceleracion' in data %}
                        <tr>
                            <td>{{ label }}</td>
                            <td>{{ "%.4f"|format(max_abs(data[comp + '_aceleracion'])) }}</td>
                        </tr>
                        {% endif %}
                        {% endfor %}
                    </table>
                </div>
                
                <div class="graph-section">
                    <h2>Gráficos de Análisis</h2>
                    <div>
                        <h3>Registro de Aceleración</h3>
                        <img src="data:image/png;base64,{{ acceleration_plot }}" style="width:100%;">
                    </div>
                    
                    {% if 'response_spectrum' in analysis_results %}
                    <div>
                        <h3>Espectro de Respuesta</h3>
                        <img src="data:image/png;base64,{{ response_spectrum_plot }}" style="width:100%;">
                    </div>
                    {% endif %}
                </div>
            </div>
        </body>
        </html>
        """
        
        # Función para convertir figura a base64
        def fig_to_base64(fig, dpi=200):
            buf = BytesIO()
            fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
            buf.seek(0)
            img_str = base64.b64encode(buf.read()).decode('utf-8')
            return img_str
        
        # Generar gráficos
        # Gráfico de aceleración
        plt.figure(figsize=(10, 6))
        for component, label in zip(['N', 'E', 'Z'], ['Norte-Sur', 'Este-Oeste', 'Vertical']):
            if f'{component}_aceleracion' in data:
                plt.plot(data['time'], data[f'{component}_aceleracion'], label=f'{label}')
        plt.title('Registro de Aceleración')
        plt.xlabel('Tiempo (s)')
        plt.ylabel('Aceleración (g)')
        plt.grid(True)
        plt.legend()
        plt.tight_layout()
        acceleration_plot = fig_to_base64(plt.gcf())
        plt.close()
        
        # Gráfico de espectro de respuesta (si existe)
        response_spectrum_plot = ""
        if 'response_spectrum' in analysis_results:
            plt.figure(figsize=(10, 6))
            plt.loglog(analysis_results['response_spectrum']['periods'], 
                      analysis_results['response_spectrum']['Sa'], 
                      label='Pseudo-aceleración')
            plt.title('Espectro de Respuesta')
            plt.xlabel('Período (s)')
            plt.ylabel('Sa (g)')
            plt.grid(True, which="both")
            plt.legend()
            plt.tight_layout()
            response_spectrum_plot = fig_to_base64(plt.gcf())
            plt.close()
        
        # Función auxiliar para obtener el valor máximo absoluto
        def max_abs(arr):
            return np.max(np.abs(arr))
        
        # Renderizar plantilla
        template = jinja2.Template(html_template)
        html_content = template.render(
            data_name=data['name'],
            generation_date=datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
            metadata=data.get('metadata', {}),
            data=data,
            max_abs=max_abs,
            analysis_results=analysis_results,
            acceleration_plot=acceleration_plot,
            response_spectrum_plot=response_spectrum_plot
        )
        
        # Guardar archivo HTML
        output_path = self.report_dir / f"{filename}.html"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Abrir en el navegador
        webbrowser.open(f'file://{output_path}')
        
        return str(output_path)
    
    def _generate_docx_report(self, data, analysis_results, filename):
        """
        Genera un reporte en formato DOCX
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("La biblioteca 'python-docx' es necesaria para generar reportes DOCX. Instálela con 'pip install python-docx'.")
        
        doc = Document()
        
        # Título
        title = doc.add_heading('Reporte de Análisis Sísmico', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Registro: {data['name']}\n").bold = True
        p.add_run(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        
        # Información del registro
        doc.add_heading('Información del Registro', level=1)
        
        # Metadatos
        if 'metadata' in data:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Parámetro'
            hdr_cells[1].text = 'Valor'
            
            for key, value in data['metadata'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
        
        # Parámetros del registro
        doc.add_heading('Parámetros del Registro', level=2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Componente'
        hdr_cells[1].text = 'PGA (g)'
        
        for component, label in zip(['N', 'E', 'Z'], ['Norte-Sur', 'Este-Oeste', 'Vertical']):
            if f'{component}_aceleracion' in data:
                acc_data = data[f'{component}_aceleracion']
                pga = np.max(np.abs(acc_data))
                
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = f"{pga:.4f}"
        
        # Gráficos
        doc.add_heading('Gráficos de Análisis', level=1)
        
        # Gráfico de aceleración
        plt.figure(figsize=(10, 6))
        for component, label in zip(['N', 'E', 'Z'], ['Norte-Sur', 'Este-Oeste', 'Vertical']):
            if f'{component}_aceleracion' in data:
                plt.plot(data['time'], data[f'{component}_aceleracion'], label=f'{label}')
        plt.title('Registro de Aceleración')
        plt.xlabel('Tiempo (s)')
        plt.ylabel('Aceleración (g)')
        plt.grid(True)
        plt.legend()
        plt.tight_layout()
        
        # Guardar temporalmente y agregar al documento
        temp_img = BytesIO()
        plt.savefig(temp_img, format='png', dpi=200, bbox_inches='tight')
        temp_img.seek(0)
        plt.close()
        
        doc.add_picture(temp_img, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Si hay resultados de espectro de respuesta
        if 'response_spectrum' in analysis_results:
            doc.add_heading('Espectro de Respuesta', level=2)
            
            # Gráfico de espectro de respuesta
            plt.figure(figsize=(10, 6))
            plt.loglog(analysis_results['response_spectrum']['periods'], 
                      analysis_results['response_spectrum']['Sa'], 
                      label='Pseudo-aceleración')
            plt.title('Espectro de Respuesta')
            plt.xlabel('Período (s)')
            plt.ylabel('Sa (g)')
            plt.grid(True, which="both")
            plt.legend()
            plt.tight_layout()
            
            # Guardar temporalmente y agregar al documento
            temp_img = BytesIO()
            plt.savefig(temp_img, format='png', dpi=200, bbox_inches='tight')
            temp_img.seek(0)
            plt.close()
            
            doc.add_picture(temp_img, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Guardar el documento
        output_path = self.report_dir / f"{filename}.docx"
        doc.save(output_path)
        
        return str(output_path)
